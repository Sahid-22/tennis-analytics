import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_style(doc):
    styles = doc.styles
    
    # Body Text Style
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.space_before = Pt(24)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.space_before = Pt(18)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.space_before = Pt(12)
    
    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_page_number(paragraph):
    p = paragraph._p
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = OxmlElement('w:r')
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)
    p.append(r)

def create_report():
    doc = Document()
    set_style(doc)
    
    # --- 1. Title Page ---
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph("Tennis Analytics:\nA Professional Sports Data Analytics Platform")
    title.style = doc.styles['Heading 1']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
        
    sub_title = doc.add_paragraph("A Project Report Submitted to")
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    uni = doc.add_paragraph("[University Name]")
    uni.style = doc.styles['Heading 2']
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    dept = doc.add_paragraph("Department of Computer Science")
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
        
    submitted_by = doc.add_paragraph("Submitted by:\n[Student Name]\n[Student ID]")
    submitted_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4):
        doc.add_paragraph()
        
    date = doc.add_paragraph("July 2026")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 2. Certificate ---
    cert_title = doc.add_heading("CERTIFICATE", level=1)
    cert_text = (
        "This is to certify that the project report entitled 'Tennis Analytics: A Professional Sports "
        "Data Analytics Platform' submitted by [Student Name] to [University Name] in partial fulfillment "
        "of the requirements for the award of the degree of Bachelor of Technology in Computer Science is "
        "a bona fide record of the project work carried out by them under my supervision and guidance.\n\n"
        "The content of the report, in full or in parts, has not been submitted to any other Institute or "
        "University for the award of any degree or diploma.\n\n\n\n\n"
        "_______________________________\n"
        "Signature of Guide/Supervisor\n"
        "Name: \n"
        "Designation: \n\n\n"
        "_______________________________\n"
        "Signature of Head of Department\n"
        "Name: "
    )
    doc.add_paragraph(cert_text)
    doc.add_page_break()

    # --- 3. Acknowledgement ---
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    ack_text = (
        "The success and final outcome of this project required a lot of guidance and assistance from many "
        "people and I am extremely privileged to have got this all along the completion of my project. All "
        "that I have done is only due to such supervision and assistance and I would not forget to thank them.\n\n"
        "I respect and thank [Guide Name], for providing me an opportunity to do the project work in [University Name] "
        "and giving us all support and guidance which made me complete the project duly. I am extremely thankful to "
        "her/him for providing such a nice support and guidance, although he had busy schedule managing the corporate affairs.\n\n"
        "I owe my deep gratitude to our project guide [Guide Name], who took keen interest on our project work and guided us "
        "all along, till the completion of our project work by providing all the necessary information for developing a good system.\n\n"
        "I would not forget to remember [HOD Name], Head of the Department of Computer Science, for her/his encouragement and "
        "more over for providing all necessary resources.\n\n"
        "Finally, I would like to express my special thanks of gratitude to my parents and friends for their continuous support "
        "and encouragement throughout my academic endeavors."
    )
    doc.add_paragraph(ack_text)
    doc.add_page_break()

    # --- 4. Abstract ---
    doc.add_heading("ABSTRACT", level=1)
    abs_text = (
        "The rapid evolution of sports analytics has created an unprecedented demand for data-driven insights in professional tennis. "
        "This project, 'Tennis Analytics', presents a comprehensive data platform designed to ingest, process, validate, and visualize "
        "complex tennis data sourced from the SportRadar API. Operating in an environment characterized by scattered API endpoints, "
        "deeply nested JSON payloads, and fluctuating data quality, the platform provides a unified and reliable data ecosystem.\n\n"
        "The core methodology involves the development of a robust ETL (Extract, Transform, Load) pipeline implemented in Python 3.11+. "
        "Data extraction employs exponential backoff and rate-limiting strategies to reliably interface with the RESTful API. "
        "The transformation phase flattens hierarchical structures and standardizes terminologies, populating a fully normalized "
        "7-table relational schema managed via SQLAlchemy. Data integrity is rigorously enforced through a bespoke Data Quality module "
        "performing over 15 automated validation checks, synthesizing a composite quality score.\n\n"
        "The user interface comprises a 7-page Streamlit dashboard featuring a premium dark theme and interactive Plotly visualizations. "
        "End-users can seamlessly navigate categorical taxonomies, competition hierarchies, player statistics, and geographical venue "
        "distributions. The platform facilitates complex analytical queries, encompassing advanced analytics like player movement leaderboards, "
        "choropleth world maps, and sunburst charts.\n\n"
        "The system's reliability is ensured through a comprehensive suite of 38 pytest automated tests and deployed via a multi-stage "
        "Docker build orchestrated by GitHub Actions. The resultant application effectively bridges the gap between raw sports data "
        "and actionable insights, providing a scalable foundation for future enhancements such as predictive modeling and real-time streaming integration."
    )
    doc.add_paragraph(abs_text)
    doc.add_page_break()

    # --- 5. Table of Contents ---
    doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_text = (
        "1. Introduction\n"
        "   1.1 Background\n"
        "   1.2 Motivation\n"
        "   1.3 Problem Statement\n"
        "   1.4 Objectives\n"
        "   1.5 Scope and Limitations\n"
        "   1.6 Organization of Report\n\n"
        "2. Literature Review\n"
        "   2.1 Sports Analytics\n"
        "   2.2 ETL Pipelines\n"
        "   2.3 Data Visualization\n"
        "   2.4 API Integration Patterns\n"
        "   2.5 Data Quality Frameworks\n"
        "   2.6 Summary of Literature\n\n"
        "3. Methodology\n"
        "   3.1 Development Methodology\n"
        "   3.2 Requirements Analysis\n"
        "   3.3 Tools and Technologies\n"
        "   3.4 Development Environment Setup\n\n"
        "4. System Design\n"
        "   4.1 System Architecture\n"
        "   4.2 Database Design\n"
        "   4.3 ETL Pipeline Design\n"
        "   4.4 Dashboard Design\n"
        "   4.5 Quality Framework Design\n\n"
        "5. Implementation\n"
        "   5.1 Project Structure\n"
        "   5.2 API Client Implementation\n"
        "   5.3 Data Transformation\n"
        "   5.4 Database Layer\n"
        "   5.5 Data Quality Module\n"
        "   5.6 Dashboard Implementation\n"
        "   5.7 Testing Implementation\n"
        "   5.8 CI/CD Pipeline\n"
        "   5.9 Docker Deployment\n\n"
        "6. Results and Discussion\n"
        "   6.1 Data Ingestion Results\n"
        "   6.2 Dashboard Functionality\n"
        "   6.3 Data Quality Results\n"
        "   6.4 Testing Results\n"
        "   6.5 Performance Analysis\n\n"
        "7. Conclusion and Future Scope\n"
        "   7.1 Summary of Work\n"
        "   7.2 Key Contributions\n"
        "   7.3 Limitations\n"
        "   7.4 Future Enhancements\n\n"
        "References\n\n"
        "Appendices\n"
    )
    doc.add_paragraph(toc_text)
    doc.add_page_break()

    # --- 6. List of Figures & 7. List of Tables ---
    doc.add_heading("LIST OF FIGURES", level=1)
    doc.add_paragraph("Figure 4.1: System Architecture Diagram")
    doc.add_paragraph("Figure 4.2: Entity Relationship (ER) Diagram")
    doc.add_paragraph("Figure 4.3: Dashboard Navigation Flow")
    doc.add_paragraph("Figure 5.1: API Client Retry Mechanism")
    doc.add_paragraph("Figure 6.1: Quality Score Gauge Visualization")
    doc.add_page_break()

    doc.add_heading("LIST OF TABLES", level=1)
    doc.add_paragraph("Table 4.1: Complete Database Schema Definition")
    doc.add_paragraph("Table 6.1: Database Ingestion Statistics")
    doc.add_paragraph("Table 6.2: Test Results Summary")
    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph("The sports analytics industry has experienced explosive growth over the past decade, transforming from rudimentary statistical tracking to sophisticated predictive modeling and real-time biometric analysis. In professional tennis, organizations such as the ATP and WTA, as well as sports betting syndicates and broad-casting networks, increasingly rely on data-driven decision making. The availability of high-fidelity data has necessitated the creation of specialized platforms capable of aggregating and analyzing this information to gain a competitive edge, evaluate player performance, and enhance fan engagement.")
    
    doc.add_heading("1.2 Motivation", level=2)
    doc.add_paragraph("Despite the abundance of tennis data, there remains a significant gap in unified platforms that seamlessly integrate data ingestion, rigorous quality validation, and interactive visualization. Existing solutions often lack transparent data pipelines or fail to present analytical results intuitively. This project was motivated by the desire to construct an end-to-end analytics platform that not only consumes raw data from a commercial provider (SportRadar API) but also transforms it into actionable insights through a reliable, automated, and scalable architecture.")
    
    doc.add_heading("1.3 Problem Statement", level=2)
    doc.add_paragraph("Working with external data APIs presents multifaceted challenges. The data retrieved is frequently encapsulated within deeply nested JSON structures, requiring complex unwrapping and normalization. Furthermore, disparate data endpoints often yield inconsistent or fragmented information, leading to referential integrity issues when persisting to a relational database. Additionally, without an automated data quality framework, end-users cannot ascertain the reliability of the insights presented. The core problem this project addresses is the seamless orchestration of an ETL pipeline that mitigates these data anomalies while feeding an interactive, high-performance dashboard.")
    
    doc.add_heading("1.4 Objectives", level=2)
    doc.add_paragraph("The primary objectives of this project are as follows:\n"
                      "1. To architect and implement a robust ETL pipeline capable of interacting with the SportRadar Tennis API, featuring exponential backoff and rate-limit handling.\n"
                      "2. To design a normalized relational database schema (comprising 7 interconnected tables) that ensures referential integrity and efficient querying.\n"
                      "3. To develop an interactive, 7-page Streamlit dashboard featuring advanced visualizations such as treemaps, sunburst charts, and choropleth maps.\n"
                      "4. To construct an automated Data Quality module that executes over 15 validation checks and generates a composite reliability score.\n"
                      "5. To ensure system reliability and maintainability through extensive automated testing (pytest) and containerized deployment (Docker).")
    
    doc.add_heading("1.5 Scope and Limitations", level=2)
    doc.add_paragraph("The scope of this project encompasses the extraction and analysis of tennis categories, competitions, complexes, venues, and competitor rankings. It leverages historical and current data provided by the SportRadar v3 API. The current iteration focuses on structured relational data and descriptive analytics. Limitations include the reliance on API rate limits enforced by the provider, which restricts the frequency of data synchronization. Additionally, the project currently implements descriptive and diagnostic analytics but does not yet incorporate predictive machine learning models for match outcome forecasting.")
    
    doc.add_heading("1.6 Organization of Report", level=2)
    doc.add_paragraph("The remainder of this report is organized as follows: Chapter 2 reviews the existing literature surrounding sports analytics, ETL pipelines, and visualization frameworks. Chapter 3 outlines the methodology and technological stack employed. Chapter 4 details the system architecture and database design. Chapter 5 discusses the implementation specifics of the data pipeline, quality framework, and dashboard. Chapter 6 presents the results and performance evaluation. Finally, Chapter 7 concludes the report and outlines avenues for future research and enhancement.")
    doc.add_page_break()

    # --- Chapter 2: Literature Review ---
    doc.add_heading("Chapter 2: Literature Review", level=1)
    
    doc.add_heading("2.1 Sports Analytics", level=2)
    doc.add_paragraph("The integration of quantitative analytics into professional sports, famously popularized by 'Moneyball' in baseball, has permeated virtually every athletic discipline, including tennis. Modern tennis analytics extends beyond simple metrics like first-serve percentages, encompassing spatial tracking data and long-term performance trend analysis. Research indicates that data-driven strategic planning significantly impacts tournament outcomes and player conditioning. However, the accessibility of unified, clean data remains a persistent barrier for independent researchers and analysts, necessitating platforms like the one proposed in this project.")

    doc.add_heading("2.2 ETL Pipelines", level=2)
    doc.add_paragraph("Extract, Transform, Load (ETL) pipelines form the backbone of modern data engineering. Best practices in ETL design emphasize idempotency, fault tolerance, and comprehensive logging. Literature on API integration stresses the importance of handling transient network failures through retry mechanisms and exponential backoff strategies to prevent cascading system failures. In the context of RESTful JSON APIs, the transformation phase is critical for flattening hierarchical documents into relational tuples suitable for structured querying.")

    doc.add_heading("2.3 Data Visualization", level=2)
    doc.add_paragraph("The transition from static reporting to interactive data visualization has been revolutionized by web-based frameworks. Traditional BI tools (e.g., Tableau, PowerBI) offer powerful capabilities but often lack the flexibility of programmatic custom applications. Frameworks like Streamlit and Dash allow data scientists to build bespoke web applications entirely in Python, bridging the gap between data processing and user interaction. Studies evaluating visualization efficacy highlight that interactive elements such as filtering, tooltips, and hierarchical zooming significantly enhance user comprehension of complex datasets.")

    doc.add_heading("2.4 API Integration Patterns", level=2)
    doc.add_paragraph("Interfacing with third-party APIs requires adherence to specific architectural patterns to ensure stability. Rate limiting—often implemented via Token Bucket or Leaky Bucket algorithms on the server side—necessitates intelligent client-side throttling. Furthermore, modern API clients utilize session pooling and asynchronous request handling to optimize throughput while remaining within permissible usage quotas. Security paradigms mandate the secure storage of API keys and the use of HTTPS for encrypted data transmission.")

    doc.add_heading("2.5 Data Quality Frameworks", level=2)
    doc.add_paragraph("Data quality is paramount in analytics; erroneous data leads to flawed conclusions. Automated validation frameworks are designed to continuously monitor data integrity across dimensions such as completeness, uniqueness, consistency, and validity. The implementation of composite scoring metrics allows stakeholders to quickly gauge the reliability of a dataset. Literature suggests that embedding quality checks directly into the ETL pipeline prevents downstream corruption and ensures that dashboards only present validated information.")

    doc.add_heading("2.6 Summary of Literature", level=2)
    doc.add_paragraph("The reviewed literature underscores the necessity of a cohesive approach to data engineering and visualization in sports analytics. While theoretical frameworks for ETL processes and data validation are well-established, their practical integration into a unified, domain-specific application presents unique engineering challenges. This project synthesizes these best practices, employing a modern Python technology stack to deliver a robust and reliable tennis analytics platform.")
    doc.add_page_break()

    # --- Chapter 3: Methodology ---
    doc.add_heading("Chapter 3: Methodology", level=1)
    
    doc.add_heading("3.1 Development Methodology", level=2)
    doc.add_paragraph("An Agile, iterative development methodology was adopted for this project. The system was developed in sequential sprints, beginning with core API connectivity, progressing through database schema design and data transformation, and culminating in the development of the frontend dashboard and CI/CD automation. This approach facilitated continuous testing and refinement, allowing for rapid adaptation to the structural nuances of the SportRadar API responses.")

    doc.add_heading("3.2 Requirements Analysis", level=2)
    doc.add_paragraph("The functional requirements of the system dictated the capability to reliably extract comprehensive tennis data, store it preserving complex relational hierarchies, and visualize it through interactive, user-driven dashboards. Non-functional requirements included high availability, fault tolerance (particularly concerning network requests), data integrity enforcement, and a modular architecture conducive to future expansion. Furthermore, the system required strict automated test coverage to ensure long-term maintainability.")

    doc.add_heading("3.3 Tools and Technologies", level=2)
    doc.add_paragraph("The project leverages a modern, Python-centric technology stack, carefully selected to align with the system requirements:\n\n"
                      "- **Python 3.11+**: The core programming language, chosen for its robust data manipulation libraries and typing enhancements.\n"
                      "- **SportRadar Tennis API v3**: The primary data source, providing authoritative, RESTful JSON data regarding global tennis events.\n"
                      "- **SQLAlchemy 2.0**: Utilized as the Object-Relational Mapper (ORM), abstracting database interactions and facilitating seamless switching between SQLite, PostgreSQL, and MySQL.\n"
                      "- **Pandas**: Employed within the transformation layer for efficient in-memory data manipulation, cleansing, and type coercion.\n"
                      "- **Streamlit**: Selected as the frontend framework to rapidly develop a highly interactive, responsive web application without requiring distinct JavaScript development.\n"
                      "- **Plotly**: Used for generating sophisticated, interactive visualizations such as choropleth maps and sunburst charts, integrated seamlessly with Streamlit.\n"
                      "- **Pytest**: The testing framework of choice, utilized to construct a comprehensive suite of 38 unit and integration tests.\n"
                      "- **Docker**: Employed to containerize the application, ensuring consistency across development and production environments via multi-stage builds.")

    doc.add_heading("3.4 Development Environment Setup", level=2)
    doc.add_paragraph("The development environment was standardized using virtual environments to isolate dependencies. Environment variables, explicitly `.env` files, were utilized to securely manage sensitive configurations such as the `SPORT_RADAR_API_KEY` and database connection strings. Version control was managed via Git, with GitHub serving as the remote repository, enabling collaborative development and the integration of automated GitHub Actions workflows for CI/CD.")
    doc.add_page_break()

    # --- Chapter 4: System Design ---
    doc.add_heading("Chapter 4: System Design", level=1)
    
    doc.add_heading("4.1 System Architecture", level=2)
    doc.add_paragraph("The system architecture follows a modular, layered design pattern comprising Data Ingestion, Data Storage, Business Logic, and Presentation layers. The pipeline initiates via the API Client, extracting raw JSON payloads which are subsequently passed to the Transformation layer. The resulting normalized Data Transfer Objects (DTOs) are validated by the Data Quality module before being persisted to the relational database via SQLAlchemy. Finally, the Streamlit dashboard queries the database to render interactive visualizations.")

    doc.add_heading("4.2 Database Design", level=2)
    doc.add_paragraph("The database schema is highly normalized, consisting of 7 distinct tables designed to eliminate data redundancy and enforce referential integrity. The core entities include `categories`, `competitions`, `complexes`, `venues`, `competitors`, and `competitor_rankings`. Foreign key relationships map competitions to their respective categories and parent competitions, and venues to their associated complexes. A dedicated `api_sync_log` table tracks data ingestion metadata, including timestamps, row counts, and payload SHA-256 checksums to ensure traceability.")

    doc.add_heading("4.3 ETL Pipeline Design", level=2)
    doc.add_paragraph("The Extract phase utilizes the `requests` library wrapped in custom retry logic utilizing exponential backoff to handle HTTP 429 (Too Many Requests) and 5xx server errors gracefully. The Transform phase involves recursive functions to flatten nested dictionaries and lists, harmonizing inconsistent country codes, and coercing data types (e.g., parsing ISO 8601 strings into datetime objects). The Load phase employs a transactional 'replace-all' strategy, clearing existing table data and bulk-inserting the new dataset to maintain a synchronized state with the external API.")

    doc.add_heading("4.4 Dashboard Design", level=2)
    doc.add_paragraph("The Streamlit dashboard is structured across 7 distinct pages, navigated via a sidebar menu. The design language employs a premium dark theme, enhancing visual contrast for the embedded Plotly charts. The pages are logically segmented into Overview KPIs, Competition Hierarchies, Competitor Statistics, Venue Geographical Distribution, Custom SQL Analysis, Advanced Analytics (Maps/Sunbursts), and Data Quality Reports. Each page incorporates dynamic filters, allowing users to slice the data by various dimensions.")

    doc.add_heading("4.5 Quality Framework Design", level=2)
    doc.add_paragraph("The Data Quality module is designed as an independent service that scrutinizes the database post-ingestion. It defines over 15 distinct rules, including Null Value Checks, Foreign Key Integrity, Unique Constraint Verification, and Range Checks (e.g., ensuring rankings are positive integers). The framework aggregates the results of these checks to compute a composite Quality Score (0-100), outputting detailed JSON and Markdown reports summarizing the health of the dataset.")

    doc.add_heading("Table 4.1: Database Schema", level=3)
    schema_table = doc.add_table(rows=1, cols=3)
    schema_table.style = 'Table Grid'
    hdr_cells = schema_table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Primary Key'
    hdr_cells[2].text = 'Key Columns & Foreign Keys'
    
    schema_data = [
        ('categories', 'category_id', 'category_name'),
        ('competitions', 'competition_id', 'competition_name, type, gender, level, parent_id, category_id (FK)'),
        ('complexes', 'complex_id', 'complex_name'),
        ('venues', 'venue_id', 'venue_name, city_name, country_name, timezone, complex_id (FK)'),
        ('competitors', 'competitor_id', 'name, country, country_code, abbreviation'),
        ('competitor_rankings', 'ID (Auto)', 'rank, movement, points, competitor_id (FK), ranking_type_id'),
        ('api_sync_log', 'ID (Auto)', 'endpoint, status_code, row_count, payload_sha256, fetched_at')
    ]
    
    for table, pk, cols in schema_data:
        row_cells = schema_table.add_row().cells
        row_cells[0].text = table
        row_cells[1].text = pk
        row_cells[2].text = cols

    doc.add_page_break()

    # --- Chapter 5: Implementation ---
    doc.add_heading("Chapter 5: Implementation", level=1)
    
    doc.add_heading("5.1 Project Structure", level=2)
    doc.add_paragraph("The repository is structured to separate concerns, isolating pipeline logic from presentation code. The `src/` directory contains sub-modules for `api/`, `db/`, `pipeline/`, and `dashboard/`. Configuration files, including Dockerfiles and workflow YAMLs, reside in the root directory. This modularity ensures that the ETL pipeline can be executed independently of the Streamlit application.")

    doc.add_heading("5.2 API Client Implementation", level=2)
    doc.add_paragraph("The `api_client.py` module defines the `SportRadarClient` class. It manages authentication headers (`x-api-key`) and utilizes the `requests.Session` object to connection pooling. A robust retry decorator is implemented, which intercepts connection timeouts and rate-limiting responses, pausing execution with an exponentially increasing delay before re-attempting the request, ensuring fault tolerance.")

    doc.add_heading("5.3 Data Transformation", level=2)
    doc.add_paragraph("The `transformers.py` module houses pure functions responsible for cleaning raw JSON. It utilizes Pandas DataFrames to perform vectorized operations where possible, such as mapping country names to standardized ISO 3166-1 alpha-3 codes. The transformation logic explicitly handles missing keys and inconsistent data types, ensuring the output rigorously adheres to the expected database schema.")

    doc.add_heading("5.4 Database Layer", level=2)
    doc.add_paragraph("Implemented using SQLAlchemy 2.0, the `models.py` file defines declarative base classes for all 7 tables. The `database.py` module manages engine creation and session lifecycle. A key implementation detail is the `replace_all_data` function, which executes within a single database transaction. It carefully orchestrates the deletion of existing records and the insertion of new records while respecting foreign key constraints, rolling back entirely if an error occurs to prevent partial updates.")

    doc.add_heading("5.5 Data Quality Module", level=2)
    doc.add_paragraph("The `quality.py` module iterates over the database tables, executing a suite of SQLAlchemy queries to identify anomalies. For instance, it verifies that no `competitor_id` in the `competitor_rankings` table lacks a corresponding entry in the `competitors` table. The module aggregates the count of passed, warned, and failed checks, applying a weighted algorithm to calculate the final composite quality score.")

    doc.add_heading("5.6 Dashboard Implementation", level=2)
    doc.add_paragraph("The Streamlit frontend utilizes custom CSS injection to achieve a cohesive dark theme. The `app.py` serves as the entry point, orchestrating navigation. Individual pages (e.g., `pages/1_Overview.py`) query the database via SQLAlchemy, convert the result sets to Pandas DataFrames, and render them using `st.dataframe` and `plotly.express`. Extensive use of Streamlit's caching mechanisms (`@st.cache_data`) minimizes database load and accelerates page rendering times.")

    doc.add_heading("5.7 Testing Implementation", level=2)
    doc.add_paragraph("A comprehensive test suite of 38 tests is implemented using `pytest`. The tests utilize mocking (via `unittest.mock`) to simulate API responses without making live HTTP calls. Fixtures are heavily employed to setup and teardown in-memory SQLite databases for testing the SQLAlchemy models and pipeline logic in isolation, achieving high code coverage.")

    doc.add_heading("5.8 CI/CD Pipeline", level=2)
    doc.add_paragraph("Continuous Integration is enforced via GitHub Actions. Upon every push or pull request to the main branch, a workflow is triggered that provisions an Ubuntu runner, sets up Python 3.11, installs dependencies, runs the `flake8` linter, and executes the `pytest` suite. This ensures that no breaking changes are merged into the production codebase.")

    doc.add_heading("5.9 Docker Deployment", level=2)
    doc.add_paragraph("The application is containerized using a multi-stage Dockerfile. The 'builder' stage installs dependencies and compiles required binaries, while the final stage copies only the necessary artifacts into a lightweight Python runtime image. A `docker-compose.yml` file is provided to orchestrate the application, exposing the Streamlit service on port 8501 and allowing optional mounting of a PostgreSQL database container.")
    doc.add_page_break()

    # --- Chapter 6: Results and Discussion ---
    doc.add_heading("Chapter 6: Results and Discussion", level=1)
    
    doc.add_heading("6.1 Data Ingestion Results", level=2)
    doc.add_paragraph("The execution of the ETL pipeline successfully ingested data across all 7 tables. Approximately 13,428 total rows were processed and persisted without error. The transactional strategy proved effective; simulated API failures triggered rollbacks, leaving the database in its previous consistent state. The `api_sync_log` accurately captured execution metadata, confirming the stability of the ingestion process.")
    
    doc.add_heading("Table 6.1: Database Statistics", level=3)
    stat_table = doc.add_table(rows=1, cols=2)
    stat_table.style = 'Table Grid'
    hdr = stat_table.rows[0].cells
    hdr[0].text = 'Table Name'
    hdr[1].text = 'Row Count'
    
    stats_data = [
        ('categories', '18'),
        ('competitions', '6,619'),
        ('complexes', '767'),
        ('venues', '4,021'),
        ('competitors', '1,000'),
        ('competitor_rankings', '1,000'),
        ('api_sync_log', '3')
    ]
    for tbl, cnt in stats_data:
        row = stat_table.add_row().cells
        row[0].text = tbl
        row[1].text = cnt

    doc.add_heading("6.2 Dashboard Functionality", level=2)
    doc.add_paragraph("The Streamlit dashboard successfully presents the data intuitively. The Overview page's animated KPI cards and the Advanced Analytics page's choropleth maps render rapidly, providing immediate geographical insights into venue distributions and competitor origins. The dynamic filtering capabilities on the Competitions and Competitors pages allow users to effortlessly drill down into specific data subsets. The custom SQL Analysis page effectively enables power users to execute arbitrary queries securely.")

    doc.add_heading("6.3 Data Quality Results", level=2)
    doc.add_paragraph("The Data Quality module execution resulted in a high composite score, reflecting the rigorous normalization performed during the transformation phase. The automated checks successfully identified minor anomalies, such as incomplete timezone data for a subset of venues, which were appropriately flagged as warnings rather than fatal errors in the detailed markdown report, proving the framework's nuanced evaluation capability.")

    doc.add_heading("6.4 Testing Results", level=2)
    doc.add_paragraph("The `pytest` suite executed successfully across all modules.")
    
    doc.add_heading("Table 6.2: Test Results Summary", level=3)
    test_table = doc.add_table(rows=1, cols=3)
    test_table.style = 'Table Grid'
    thdr = test_table.rows[0].cells
    thdr[0].text = 'Test Category'
    thdr[1].text = 'Number of Tests'
    thdr[2].text = 'Pass Rate'
    
    test_data = [
        ('API Client Mock Tests', '8', '100%'),
        ('Data Transformation', '12', '100%'),
        ('Database Models & ORM', '10', '100%'),
        ('Quality Framework', '8', '100%')
    ]
    for cat, num, rate in test_data:
        row = test_table.add_row().cells
        row[0].text = cat
        row[1].text = num
        row[2].text = rate

    doc.add_heading("6.5 Performance Analysis", level=2)
    doc.add_paragraph("Performance profiling indicated that the Streamlit caching effectively mitigated database load; after the initial query, subsequent page reloads rendered in under 100 milliseconds. The ETL pipeline execution time is dominated by API network latency and rate limiting, rather than database insertion, validating the efficiency of the SQLAlchemy bulk insert mechanisms.")
    doc.add_page_break()

    # --- Chapter 7: Conclusion and Future Scope ---
    doc.add_heading("Chapter 7: Conclusion and Future Scope", level=1)
    
    doc.add_heading("7.1 Summary of Work", level=2)
    doc.add_paragraph("This project successfully delivered a professional-grade Tennis Analytics platform. By orchestrating a robust ETL pipeline, a highly normalized database schema, and an interactive Streamlit dashboard, the system effectively transforms raw JSON payloads from the SportRadar API into an accessible, actionable intelligence ecosystem. The integration of a dedicated Data Quality module and comprehensive automated testing ensures the platform remains reliable and maintainable.")

    doc.add_heading("7.2 Key Contributions", level=2)
    doc.add_paragraph("The key technical contributions of this project include the implementation of resilient API extraction logic utilizing exponential backoff, the design of an automated data validation framework that generates composite quality scores, and the creation of an advanced analytical dashboard utilizing seamless integration between Streamlit and Plotly for geospatial and hierarchical data visualization.")

    doc.add_heading("7.3 Limitations", level=2)
    doc.add_paragraph("Current limitations primarily involve the static nature of the data updates; the system requires a manual or scheduled trigger to execute the synchronization pipeline. Furthermore, the reliance on third-party API rate limits restricts high-frequency data polling, making real-time, in-match analytics unfeasible with the current subscription tier.")

    doc.add_heading("7.4 Future Enhancements", level=2)
    doc.add_paragraph("Future iterations of the platform could incorporate historical trend analysis by implementing temporal database schemas, tracking player ranking movements over multi-year periods. The integration of Machine Learning (ML) models utilizing scikit-learn or TensorFlow could enable predictive analytics, forecasting match outcomes based on historical performance metrics. Additionally, extending the ETL pipeline to support real-time WebSocket streams would allow for live dashboard updates during tournament play.")
    doc.add_page_break()

    # --- References ---
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "[1] SportRadar, 'Tennis API v3 Documentation', SportRadar AG, 2024. [Online]. Available: https://developer.sportradar.com/",
        "[2] Streamlit Inc., 'Streamlit Documentation: The fastest way to build data apps', 2024. [Online]. Available: https://docs.streamlit.io/",
        "[3] M. Bayer, 'SQLAlchemy 2.0 Documentation', 2024. [Online]. Available: https://docs.sqlalchemy.org/",
        "[4] Python Software Foundation, 'Python 3.11 Documentation', 2024. [Online]. Available: https://docs.python.org/3.11/",
        "[5] Plotly Technologies Inc., 'Plotly Python Open Source Graphing Library', 2024. [Online]. Available: https://plotly.com/python/",
        "[6] Docker Inc., 'Docker Documentation', 2024. [Online]. Available: https://docs.docker.com/",
        "[7] W. McKinney, 'Data Structures for Statistical Computing in Python', Proceedings of the 9th Python in Science Conference, 2010.",
        "[8] R. Kimball and M. Ross, 'The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling', 3rd ed. Wiley, 2013.",
        "[9] T. H. Davenport, 'Analytics in Sports: The New Science of Winning', International Institute for Analytics, 2014.",
        "[10] M. Kleppmann, 'Designing Data-Intensive Applications: The Big Ideas Behind Reliable, Scalable, and Maintainable Systems', O'Reilly Media, 2017.",
        "[11] A. Beaulieu, 'Learning SQL', 3rd ed. O'Reilly Media, 2020.",
        "[12] J. VanderPlas, 'Python Data Science Handbook: Essential Tools for Working with Data', O'Reilly Media, 2016.",
        "[13] L. Bass, P. Clements, and R. Kazman, 'Software Architecture in Practice', 3rd ed. Addison-Wesley Professional, 2012.",
        "[14] H. Garcia-Molina, J. D. Ullman, and J. Widom, 'Database Systems: The Complete Book', 2nd ed. Pearson Prentice Hall, 2008.",
        "[15] K. Beck, 'Test-Driven Development: By Example', Addison-Wesley Professional, 2002."
    ]
    for ref in refs:
        doc.add_paragraph(ref)
    doc.add_page_break()

    # --- Appendices ---
    doc.add_heading("APPENDICES", level=1)
    
    doc.add_heading("Appendix A: Project File Structure", level=2)
    doc.add_paragraph("SportRadar_Tennis_Analytics/\n├── .github/workflows/ci.yml\n├── src/\n│   ├── api/api_client.py\n│   ├── db/models.py\n│   ├── db/database.py\n│   ├── pipeline/transformers.py\n│   ├── pipeline/quality.py\n│   ├── dashboard/app.py\n│   └── dashboard/pages/\n├── tests/\n├── requirements.txt\n├── Dockerfile\n└── generate_report.py")
    
    doc.add_heading("Appendix B: SQL Schema DDL", level=2)
    doc.add_paragraph("CREATE TABLE categories (\n    category_id VARCHAR PRIMARY KEY,\n    category_name VARCHAR NOT NULL\n);")
    
    doc.add_heading("Appendix C: Sample API Response JSON", level=2)
    doc.add_paragraph('{\n  "competitor": {\n    "id": "sr:competitor:123",\n    "name": "Federer, Roger",\n    "country": "Switzerland",\n    "abbreviation": "FED"\n  }\n}')
    
    doc.add_heading("Appendix D: Quality Check Definitions", level=2)
    doc.add_paragraph("Check 1: Non-null primary keys\nCheck 2: Foreign key validity in competitor_rankings\nCheck 3: Positive integer validation for rank and points")

    doc.save(r"d:\Labmentix Project\SportRadar_Tennis_Analytics\Tennis_Analytics_Project_Report.docx")
    print("Report generated successfully.")

if __name__ == '__main__':
    create_report()
